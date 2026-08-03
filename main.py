import os
import re
import io
import sqlite3
from fastapi import FastAPI, HTTPException, UploadFile, File, Header
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from openai import OpenAI
import docx
import pypdf
from dotenv import load_dotenv

# تحميل المتغيرات السرية
load_dotenv()

app = FastAPI(title="Arabic AI Humanizer SaaS")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "static")
DB_PATH = os.path.join(BASE_DIR, "users_data.db")

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "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")


# --- 🗄️ تهيئة قاعدة البيانات ---
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            user_id TEXT PRIMARY KEY,
            plan_type TEXT DEFAULT 'free',  -- free, monthly, pro
            attempts_used INTEGER DEFAULT 0,
            words_used INTEGER DEFAULT 0,
            words_limit INTEGER DEFAULT 0
        )
    ''')
    conn.commit()
    conn.close()


init_db()


def get_or_create_user(user_id: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT user_id, plan_type, attempts_used, words_used, words_limit FROM users WHERE user_id = ?",
                   (user_id,))
    user = cursor.fetchone()

    if not user:
        cursor.execute(
            "INSERT INTO users (user_id, plan_type, attempts_used, words_used, words_limit) VALUES (?, 'free', 0, 0, 0)",
            (user_id,))
        conn.commit()
        user = (user_id, 'free', 0, 0, 0)

    conn.close()
    return {
        "user_id": user[0],
        "plan_type": user[1],
        "attempts_used": user[2],
        "words_used": user[3],
        "words_limit": user[4]
    }


def update_user_usage(user_id: str, words_added: int):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        UPDATE users 
        SET attempts_used = attempts_used + 1, words_used = words_used + ? 
        WHERE user_id = ?
    ''', (words_added, user_id))
    conn.commit()
    conn.close()


# --- ⚙️ نماذج البيانات ---
class HumanizeRequest(BaseModel):
    text: str
    style: str = "standard"
    length: str = "same"
    user_id: str = "guest_user"


class ExportRequest(BaseModel):
    text: str


def clean_output_text(text: str) -> str:
    pattern = r'[^\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF0-9\s.,!?::؛"\'-()،]'
    text = re.sub(pattern, '', text)
    text = re.sub(r'[*#_`~{}]', '', text)
    text = re.sub(r'(?<=\s)[بجحخدذرزسشصضطظعغفقكلمنهي](?=\s)', '', text)
    text = re.sub(r'^\s*[,،.]\s*', '', text)
    return re.sub(r' +', ' ', text).strip()


@app.get("/")
async def read_index():
    return FileResponse(os.path.join(STATIC_DIR, "index.html"))


# 📊 1. معرفة بيانات واستهلاك المستخدم الحالي
@app.get("/api/user-status/{user_id}")
async def get_user_status(user_id: str):
    user = get_or_create_user(user_id)
    return {
        "status": "success",
        "user": user
    }


# 📄 2. قراءة المستندات
@app.post("/api/extract-file")
async def extract_file(file: UploadFile = File(...)):
    filename = file.filename.lower()
    content = await file.read()
    extracted_text = ""

    try:
        if filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif filename.endswith(".pdf"):
            reader = pypdf.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    extracted_text += t + "\n"
        elif filename.endswith(".txt"):
            extracted_text = content.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(status_code=400, detail="نوع الملف غير مدعوم")

        return {"status": "success", "text": extracted_text.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"حدث خطأ أثناء قراءة الملف: {str(e)}")


# 📥 3. تصدير Word
@app.post("/api/export-docx")
async def export_docx(req: ExportRequest):
    if not req.text or not req.text.strip():
        raise HTTPException(status_code=400, detail="لا يوجد نص لتصديره")

    doc = docx.Document()
    for paragraph_text in req.text.split("\n"):
        if paragraph_text.strip():
            p = doc.add_paragraph(paragraph_text.strip())
            p.paragraph_format.bidi = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Monasseq_Document.docx"}
    )


# ⚡ 4. المعالجة المباشرة مع التحقق من الباقة والمحاولات
@app.post("/api/humanize-stream")
async def humanize_stream(req: HumanizeRequest):
    if not req.text or not req.text.strip():
        raise HTTPException(status_code=400, detail="النص لا يمكن أن يكون فارغاً")

    # فحص حساب واستهلاك المستخدم
    user = get_or_create_user(req.user_id)

    # 🚫 تطبيق شرط المحاولات الـ 3 للخطة المجانية
    if user["plan_type"] == "free" and user["attempts_used"] >= 3:
        raise HTTPException(
            status_code=403,
            detail="LIMIT_REACHED: لقد استهلكت المحاولات المجانية الـ 3! يُرجى الاشتراك لمتابعة الاستخدام."
        )

    # 🚫 تطبيق شرط باقة الكلمات المحدودة (Monthly)
    if user["plan_type"] == "monthly" and user["words_used"] >= user["words_limit"]:
        raise HTTPException(
            status_code=403,
            detail="LIMIT_REACHED: لقد استنفدت رصيد الكلمات المتاح بباقتك الشهرية! يُرجى الترقية إلى باقة Pro."
        )

    style_prompts = {
        "standard": "صحفي سلس وممتع للقراءة بأسلوب كاتب عربي محترف.",
        "creative": "أدبي إبداعي وبليغ يعتمد على جماليات اللغة العربية وتدفق الجمل.",
        "academic": "رصين ومباشر بلغة الأبحاث والدراسات الموثقة دون حشو.",
        "professional": "عملي واحترافي موجه لبيئة الأعمال والشركات."
    }

    length_prompts = {
        "same": "حافظ على نفس حجم وطول النص الأصلي تقريباً دون اختصار شديد أو مط زائد.",
        "short": "اقتضب واختصر النص ليكون مكثفاً ومباشراً وبأقل عدد كلمات دون الإخلال بالمضمون.",
        "detailed": "وسع الشرح والأفكار بتفاصيل وأمثلة إضافية تجعل النص مفصلاً وشاملاً."
    }

    selected_style = style_prompts.get(req.style, style_prompts["standard"])
    selected_length = length_prompts.get(req.length, length_prompts["same"])

    system_prompt = f"""أنت كاتب ومحرر لغوي عربي محترف ومترسخ في البلاغة. وظيفتك إعادة كتابة النص ليبدو مكتوباً بقلم بشري طبيعي وسلس.

التعليمات الصارمة:
1. العربية الخالصة والإملاء الدقيق: اكتب بلغة عربية فصحى سليمة ومضبوطة إملائياً 100%. يُمنع كتابة أي أحرف أجنبية أو ترك أحرف مفردة تائهة.
2. التحكم بالطول: {selected_length}
3. حذف الكليشيهات تماماً: استبعد تماماً عبارات الـ AI الشهيرة مثل (حجر الزاوية، مما لا شك فيه، في الآونة الأخيرة، في عصرنا الحالي، سلاح ذو حدين، يجدر بالذكر، ختاماً، في ختام المطاف).
4. الروابط البلاغية الطبيعية: استخدم أساليب الربط العربية البشرية السلسة للتنقل بين الأفكار.
5. الإخراج النظيف: اكتب النص المُعدل مباشرة دون مقدمات أو تحيات أو رموز ماركداون.

الأسلوب المطلوب: {selected_style}"""

    # تسجيل الاستهلاك
    words_count = len(req.text.split())
    update_user_usage(req.user_id, words_count)

    def generate_text_chunks():
        try:
            client = OpenAI(
                api_key=GROQ_API_KEY,
                base_url="https://api.groq.com/openai/v1"
            )
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"أعد صياغة هذا النص بأسلوب بشري طبيعي وسليم:\n\n{req.text}"}
                ],
                temperature=0.4,
                stream=True
            )
            for chunk in response:
                if chunk.choices and chunk.choices[0].delta.content:
                    yield chunk.choices[0].delta.content
        except Exception as e:
            yield f"\n[حدث خطأ أثناء الاتصال: {str(e)}]"

    return StreamingResponse(generate_text_chunks(), media_type="text/plain")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8001)